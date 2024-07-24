# -*- coding: utf-8 -*-
"""
Created on Sat Jun 15 08:30:54 2024

@author: sebastian
"""

# import o
from pathlib import Path
import requests
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from bs4 import BeautifulSoup
from PyPDF2 import PdfWriter
from fpdf import FPDF
from PIL import Image
import pandas as pd
import time
import openpyxl
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches

def write_text(fn: str, txt: str):
    with open(fn, 'w') as file:
        file.write(txt)

def create_pdf(fname, image_files):
    pdf_writer = PdfWriter()
    pdf_path = f"{fname}.pdf"
    for image_file in image_files:
        image = Image.open(image_file)
        pdf_bytes = image.convert('RGB')
        pdf_bytes.save("temp_image.pdf")
        pdf_reader = PdfWriter()
        pdf_reader.append(open("temp_image.pdf", "rb"))
        pdf_writer.add_page(pdf_reader.pages[0])
        Path("temp_image.pdf").unlink(missing_ok=True)
    with open(pdf_path, "wb") as out_file:
        pdf_writer.write(out_file)
    print(f"PDF '{pdf_path}' created successfully.")

def read_hyperlinks_from_excel(file_path, sheet_name, column_name):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheet = wb[sheet_name]
    df = pd.DataFrame(sheet.values)
    cols = df.iloc[0]
    col_idx = cols.index[cols.eq(column_name)].values[0]
    links = []
    for row in range(1, sheet.max_row + 1):
        t = sheet.cell(row=row + 1, column=col_idx + 1)
        if t.hyperlink:
            links.append(t.hyperlink.target)
    return links

def init_webdriver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_service = Service("path/to/chromedriver")  # specify your chromedriver path
    driver = webdriver.Chrome(service=chrome_service, options=chrome_options)
    return driver

def get_product_details(driver, url):
    driver.maximize_window()
    driver.get(url)
    time.sleep(5)
    content = driver.page_source.encode('utf-8').strip()
    soup = BeautifulSoup(content, 'html.parser')

    # Get product name
    product_name = soup.find('span', {'id': 'productTitle'}).get_text(strip=True)
    
    
    image_container = soup.find('div', {'class': 'imgTagWrapper', 'id': 'imgTagWrapperId'})
    # Get product images
    # image_divs = image_container.find_all('img', {'alt':product_name})

    # image_divs = image_container.find_all('img', {'class':"a-dynamic-image"})
    # image_divs = soup.find_all('img', {'class': 'imgTagWrapper'})
    # image_urls = [img['data-a-dynamic-image'] for img in image_divs]
    
    # image_urls = image_container.img.get('data-a-dynamic-image')#.keys()
    # image_urls = list(json.loads(image_urls).keys())
    
    image_urls = []
    for i in driver.find_elements(By.CSS_SELECTOR, '#altImages .imageThumbnail'):
      hover = ActionChains(driver).move_to_element(i)
      hover.perform()
      image_urls.append(driver.find_element(By.CSS_SELECTOR,'.image.item.maintain-height.selected img').get_attribute('src'))

    return product_name, image_urls

def download_images(product_name, image_urls):
    path = product_name
    path.mkdir(parents = True, exist_ok = True)
    for idx, img_url in enumerate(image_urls):
        img_data = requests.get(img_url).content
        # with open(path / f'image_{idx + 1}.jpg', 'wb') as handler:
        with (path / f'image_{idx + 1}.jpg').open(mode = 'wb') as handler:
            handler.write(img_data)
            
def get_images(image_urls) -> list:
    imgs = []
    for idx, img_url in enumerate(image_urls):
        img_data = requests.get(img_url).content
        imgs.append(img_data)
    return imgs
        


    
def create_pdf_with_heading(pdf, product_name):
    pdf.add_page()
    pdf.set_font("Arial", size=16)
    pdf.cell(200, 10, txt=product_name, ln=True, align='C')
    
def create_doc_with_heading(doc, product_name):
    doc.add_heading(product_name, level=1)

def add_images_to_pdf(pdf, imgs):
    for img_data in imgs:
        img = Image.open(BytesIO(img_data))
        img_rgb = img.convert('RGB')
        img_buffer = BytesIO()
        img_rgb.save(img_buffer, format='JPEG')
        pdf.add_page()
        pdf.image(img_buffer, x=10, y=20, w=180)

def add_images_to_doc(doc, imgs):
    for img_data in imgs:
        img = Image.open(BytesIO(img_data))
        img_rgb = img.convert('RGB')
        img_buffer = BytesIO()
        img_rgb.save(img_buffer, format='JPEG')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))

def make_pdf(amazon_links: list, fname: str):
    driver = webdriver.Chrome()
    pdf = FPDF()
    for link in amazon_links:
        try:
            product_full_name, image_urls = get_product_details(driver, link)
            create_pdf_with_heading(pdf, product_full_name)
            if image_urls:
                imgs = get_images(image_urls)
                add_images_to_pdf(pdf, imgs)
                # print(f"Images for '{product_full_name.split(',')[0]}' added to PDF successfully.")
            else:
                print(f"No images found for '{product_full_name.split(',')[0]}'.")
        except Exception as e:
            print(f'Failed to make pdf for {link}: {str(e)}')
    driver.quit()
    pdf.output(fname)
    print(f"Combined PDF {fname} created successfully.")
    
def make_doc(amazon_links: list, fname: str):
    driver = webdriver.Chrome()
    doc = Document()
    for link in amazon_links:
        try:
            product_full_name, image_urls = get_product_details(driver, link)
            create_doc_with_heading(doc, product_full_name)
            if image_urls:
                imgs = get_images(image_urls)
                add_images_to_doc(doc, imgs)
                # print(f"Images for '{product_full_name.split(',')[0]}' added to PDF successfully.")
            
            else:
                print(f"No images found for '{product_full_name.split(',')[0]}'.")
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_text(f'Link to this product:\n{link}')
            # p.add_hyperlink('Link to this product', target=link)
        except Exception as e:
            print(f'Failed to make pdf for {link}: {str(e)}')
    driver.quit()
    doc.save(fname)
    print(f"Combined document {fname} created successfully.")
            
def download_to_folders(amazon_links: list, folder = 'product_images'):
    # driver = init_webdriver()
    driver = webdriver.Chrome()
    for link in amazon_links:
        try:
            product_full_name, image_urls = get_product_details(driver, link)
            product_name = product_full_name.split(',')[0]
            product_name = product_name.split('/')[0]
            target_folder = Path(f'{folder}/{product_name}')
            download_images(target_folder, image_urls)
            file_str = f'PRODUCT NAME: {product_full_name}\n\n\nLINK: {link}'
            with (target_folder/ 'README.txt').open('w', encoding ="utf-8") as file:
                file.write(file_str)
                
            print(f"Images for '{product_name}' downloaded successfully.")
        except Exception as e:
            print(f"Failed to download images for {link}: {str(e)}")
    driver.quit()
    
if __name__ == "__main__":
    
    # bestsellers = pd.read_excel('amazon_bestseller.xlsx')
    
    amazon_links = read_hyperlinks_from_excel('amazon_bestseller.xlsx', 'Sheet1', 'Link') #bestsellers['Link'].to_list()
    # download_to_folders(amazon_links)
    make_doc(amazon_links, 'products.docx')
